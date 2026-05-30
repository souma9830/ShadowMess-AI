"""
ShadowMesh — Task 11.5: Decoy Document Generator
=================================================
Generates convincing Office documents (DOCX), PDFs, and plain-text files
with embedded canary token URLs.  When an attacker exfiltrates a document
and later opens it, the embedded URL fires the existing canary pipeline,
providing post-exfiltration intelligence.

Architecture
------------
  DocGenerator.generate_for_node(node_id, node_type, base_url)
      ├─ picks templates appropriate for the node type
      ├─ creates a CanaryToken via canary_manager
      ├─ renders DOCX / PDF / TXT bytes with the canary URL embedded
      ├─ stores (bytes, filename, mime, created_at) in _store[token_id]
      └─ returns List[DocMeta]

  GET /api/docs/{token_id}/{filename}
      ├─ validates token_id exists in canary_manager
      ├─ validates filename matches stored filename
      ├─ streams bytes with correct Content-Type
      └─ fires a data_access callback (document downloaded = intelligence)

Storage
-------
  In-memory dict bounded by _MAX_DOCS (2000).  Oldest entry evicted when
  full.  Documents older than _DOC_TTL_S (24 h) are evicted on access.
  Designed for future migration to Redis (store bytes as hash field) or
  S3 (store key = token_id, value = object).

Security invariants
-------------------
  - All embedded "credentials" use FAKEAKIA / SHADOWMESH-FAKE markers.
  - filename parameter is validated against stored value — never used as
    a filesystem path.
  - No real secrets, no real employee data.
"""

from __future__ import annotations

import io
import logging
import time
import uuid
from dataclasses import dataclass
from typing import Dict, List, Optional, Tuple

log = logging.getLogger("document_generator")

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

_MAX_DOCS   = 2000          # hard cap on in-memory document store
_DOC_TTL_S  = 86_400        # 24 h — documents older than this are evicted on access

# MIME types
_MIME: Dict[str, str] = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf":  "application/pdf",
    "txt":  "text/plain; charset=utf-8",
}

# ---------------------------------------------------------------------------
# Document templates — (filename_template, format, node_types)
# ---------------------------------------------------------------------------

_TEMPLATES: List[Dict] = [
    # Finance / executive documents
    {
        "filename":   "Q2_2025_Financial_Report.pdf",
        "fmt":        "pdf",
        "title":      "Q2 2025 Financial Report",
        "subject":    "Quarterly financial summary — CONFIDENTIAL",
        "node_types": {"file_server", "smb_server", "web_server", "api_gateway"},
    },
    {
        "filename":   "Payroll_May2025.xlsx.docx",
        "fmt":        "docx",
        "title":      "Payroll — May 2025",
        "subject":    "Monthly payroll register — HR CONFIDENTIAL",
        "node_types": {"file_server", "smb_server", "workstation"},
    },
    {
        "filename":   "Employee_Directory_2025.docx",
        "fmt":        "docx",
        "title":      "Employee Directory 2025",
        "subject":    "Internal staff directory — DO NOT DISTRIBUTE",
        "node_types": {"file_server", "smb_server", "auth_service"},
    },
    {
        "filename":   "VPN_Credentials_Backup.txt",
        "fmt":        "txt",
        "title":      "VPN Credentials Backup",
        "subject":    "Emergency access credentials — RESTRICTED",
        "node_types": {"file_server", "smb_server", "auth_service", "workstation"},
    },
    {
        "filename":   "AWS_Access_Keys_Prod.txt",
        "fmt":        "txt",
        "title":      "AWS Production Access Keys",
        "subject":    "Cloud infrastructure credentials — TOP SECRET",
        "node_types": {"api_gateway", "web_server", "db_server"},
    },
    {
        "filename":   "Database_Credentials.txt",
        "fmt":        "txt",
        "title":      "Database Credentials",
        "subject":    "Production database access — CONFIDENTIAL",
        "node_types": {"db_server", "api_gateway", "auth_service"},
    },
    {
        "filename":   "Network_Architecture_2025.pdf",
        "fmt":        "pdf",
        "title":      "Network Architecture Overview 2025",
        "subject":    "Internal network topology — RESTRICTED",
        "node_types": {"web_server", "api_gateway", "file_server"},
    },
    {
        "filename":   "Incident_Response_Runbook.docx",
        "fmt":        "docx",
        "title":      "Incident Response Runbook",
        "subject":    "Security operations procedures — INTERNAL",
        "node_types": {"auth_service", "web_server", "api_gateway"},
    },
]

# Default templates used when node_type doesn't match any specific set
_DEFAULT_TEMPLATES = [t for t in _TEMPLATES if "file_server" in t["node_types"]]


# ---------------------------------------------------------------------------
# Storage entry
# ---------------------------------------------------------------------------

@dataclass
class _DocEntry:
    token_id:   str
    node_id:    str
    filename:   str
    mime:       str
    data:       bytes
    created_at: float


@dataclass
class DocMeta:
    """Public metadata returned to callers after document generation."""
    token_id:  str
    node_id:   str
    filename:  str
    fmt:       str
    url:       str       # relative URL: /api/docs/{token_id}/{filename}
    canary_url: str      # relative URL: /api/canary/{token_id}


# ---------------------------------------------------------------------------
# Bounded in-memory store
# ---------------------------------------------------------------------------

class _DocStore:
    """
    Bounded dict[token_id → _DocEntry] with TTL eviction on access.
    Thread-safe for single-process use (GIL protects dict operations).
    """

    def __init__(self, max_size: int = _MAX_DOCS, ttl: float = _DOC_TTL_S) -> None:
        self._store: Dict[str, _DocEntry] = {}
        self._max   = max_size
        self._ttl   = ttl

    def put(self, entry: _DocEntry) -> None:
        if len(self._store) >= self._max:
            # Evict oldest (dict preserves insertion order in Python 3.7+)
            oldest_key = next(iter(self._store))
            del self._store[oldest_key]
            log.debug("[store] Evicted oldest document (cap=%d)", self._max)
        self._store[entry.token_id] = entry

    def get(self, token_id: str) -> Optional[_DocEntry]:
        entry = self._store.get(token_id)
        if entry is None:
            return None
        if time.time() - entry.created_at > self._ttl:
            del self._store[token_id]
            log.debug("[store] Evicted expired document %s", token_id)
            return None
        return entry

    def remove_for_node(self, node_id: str) -> int:
        keys = [k for k, v in self._store.items() if v.node_id == node_id]
        for k in keys:
            del self._store[k]
        return len(keys)

    def count(self) -> int:
        return len(self._store)


# ---------------------------------------------------------------------------
# Renderers
# ---------------------------------------------------------------------------

def _render_txt(title: str, canary_url: str, full_canary_url: str) -> bytes:
    """
    Render a plain-text decoy document.
    The canary URL is embedded as a "verification link" — plausible in
    credential/config files that reference internal portals.
    """
    content = (
        f"{title}\n"
        f"{'=' * len(title)}\n\n"
        f"Classification: CONFIDENTIAL — Internal Use Only\n"
        f"Generated: {time.strftime('%Y-%m-%d %H:%M UTC', time.gmtime())}\n\n"
    )

    if "VPN" in title or "Credentials" in title or "Keys" in title:
        content += (
            "Access Credentials\n"
            "------------------\n"
            "Primary:   FAKEAKIA_shadowmesh_prod_access_key\n"
            "Secondary: SHADOWMESH-FAKE-secret-do-not-use\n"
            "Region:    us-east-1\n\n"
            "Verification Portal (required for activation):\n"
            f"  {full_canary_url}\n\n"
            "NOTE: Credentials expire 90 days after issue.\n"
            "Contact IT Security to renew: security@corp.shadowmesh.internal\n"
        )
    else:
        content += (
            "Document Contents\n"
            "-----------------\n"
            "This document contains sensitive corporate information.\n"
            "Unauthorized distribution is prohibited.\n\n"
            "To verify document authenticity, visit:\n"
            f"  {full_canary_url}\n\n"
            "For questions contact: compliance@corp.shadowmesh.internal\n"
        )

    return content.encode("utf-8")


def _render_docx(title: str, subject: str, canary_url: str, full_canary_url: str) -> bytes:
    """
    Render a DOCX decoy document using python-docx.
    Embeds the canary URL as a hyperlink in the document body.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import lxml.etree as etree

    doc = Document()

    # Core properties (metadata)
    core = doc.core_properties
    core.title   = title
    core.subject = subject
    core.author  = "Corporate IT"
    core.company = "ShadowMesh Corp"
    core.keywords = "confidential internal"

    # Title heading
    heading = doc.add_heading(title, level=1)
    heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # Metadata paragraph
    meta = doc.add_paragraph()
    meta.add_run(f"Classification: ").bold = True
    meta.add_run("CONFIDENTIAL — Internal Use Only\n")
    meta.add_run(f"Date: ").bold = True
    meta.add_run(time.strftime("%B %d, %Y", time.gmtime()) + "\n")
    meta.add_run(f"Department: ").bold = True
    meta.add_run("Corporate Finance & Operations\n")

    doc.add_paragraph()  # spacer

    # Body content
    doc.add_heading("Summary", level=2)
    doc.add_paragraph(
        "This document contains sensitive corporate information and is intended "
        "solely for authorized personnel. Unauthorized access, copying, or "
        "distribution is strictly prohibited and may result in legal action."
    )

    doc.add_heading("Contents", level=2)
    for item in ["Executive Summary", "Financial Data", "Personnel Records",
                 "Access Credentials", "Appendix"]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()

    # Canary hyperlink paragraph
    doc.add_heading("Document Verification", level=2)
    p = doc.add_paragraph("To verify the authenticity of this document, click: ")

    # Build hyperlink XML element
    r_id = doc.part.relate_to(
        full_canary_url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(u)
    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.text = "Verify Document"
    new_run.append(t)
    hyperlink.append(new_run)
    p._p.append(hyperlink)

    doc.add_paragraph()
    doc.add_paragraph(
        "If the link above does not work, paste the following URL into your browser:\n"
        f"{full_canary_url}"
    ).runs[0].font.size = Pt(9)

    # Footer note
    doc.add_paragraph()
    footer_p = doc.add_paragraph(
        "CONFIDENTIAL — This document is the property of ShadowMesh Corp. "
        "All rights reserved. Retention period: 7 years."
    )
    footer_p.runs[0].font.size = Pt(8)
    footer_p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_pdf(title: str, subject: str, canary_url: str, full_canary_url: str) -> bytes:
    """
    Render a PDF decoy document using reportlab.
    Embeds the canary URL as a clickable link.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.platypus import Table, TableStyle

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        rightMargin=2.5 * cm,
        leftMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
        title=title,
        subject=subject,
        author="Corporate IT",
        creator="ShadowMesh Document Management System v3.2",
    )

    styles = getSampleStyleSheet()
    style_title = ParagraphStyle(
        "DocTitle",
        parent=styles["Title"],
        fontSize=20,
        textColor=colors.HexColor("#1F497D"),
        spaceAfter=6,
    )
    style_h2 = ParagraphStyle(
        "H2",
        parent=styles["Heading2"],
        fontSize=13,
        textColor=colors.HexColor("#2E74B5"),
        spaceBefore=12,
        spaceAfter=4,
    )
    style_body = styles["BodyText"]
    style_small = ParagraphStyle(
        "Small",
        parent=styles["BodyText"],
        fontSize=8,
        textColor=colors.grey,
    )
    style_link = ParagraphStyle(
        "Link",
        parent=styles["BodyText"],
        textColor=colors.HexColor("#0563C1"),
    )

    story = [
        Paragraph(title, style_title),
        HRFlowable(width="100%", thickness=1, color=colors.HexColor("#2E74B5")),
        Spacer(1, 0.3 * cm),
        # Metadata table
        Table(
            [
                ["Classification:", "CONFIDENTIAL — Internal Use Only"],
                ["Date:", time.strftime("%B %d, %Y", time.gmtime())],
                ["Department:", "Corporate Finance & Operations"],
                ["Document ID:", f"DOC-{uuid.uuid4().hex[:8].upper()}"],
            ],
            colWidths=[4 * cm, None],
            style=TableStyle([
                ("FONTNAME",  (0, 0), (0, -1), "Helvetica-Bold"),
                ("FONTSIZE",  (0, 0), (-1, -1), 10),
                ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#333333")),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]),
        ),
        Spacer(1, 0.5 * cm),
        Paragraph("Executive Summary", style_h2),
        Paragraph(
            "This document contains sensitive corporate information and is intended "
            "solely for authorized personnel. Unauthorized access, copying, or "
            "distribution is strictly prohibited and may result in legal action "
            "under applicable data protection regulations.",
            style_body,
        ),
        Spacer(1, 0.3 * cm),
        Paragraph("Contents", style_h2),
        Paragraph("• Executive Summary<br/>• Financial Data<br/>• Personnel Records<br/>"
                  "• Access Credentials<br/>• Appendix", style_body),
        Spacer(1, 0.5 * cm),
        Paragraph("Document Verification", style_h2),
        Paragraph(
            "To verify the authenticity of this document, click the link below or "
            "paste it into your browser:",
            style_body,
        ),
        Spacer(1, 0.2 * cm),
        Paragraph(
            f'<a href="{full_canary_url}" color="#0563C1">Verify Document Authenticity</a>',
            style_link,
        ),
        Spacer(1, 0.2 * cm),
        Paragraph(full_canary_url, style_small),
        Spacer(1, 1 * cm),
        HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey),
        Spacer(1, 0.2 * cm),
        Paragraph(
            "CONFIDENTIAL — This document is the property of ShadowMesh Corp. "
            "All rights reserved. Retention period: 7 years. "
            "Document management system: DMS v3.2",
            style_small,
        ),
    ]

    doc.build(story)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# DocGenerator
# ---------------------------------------------------------------------------

class DocGenerator:
    """
    Singleton document generator.

    Generates decoy documents for a given node, stores them in the bounded
    in-memory store, and returns metadata for integration with the SMB/HTTP
    honeypots and the FastAPI download route.
    """

    def __init__(self) -> None:
        self._store = _DocStore()

    def generate_for_node(
        self,
        node_id: str,
        node_type: str,
        base_url: str,
        count: int = 2,
    ) -> List[DocMeta]:
        """
        Generate `count` decoy documents for `node_id`.

        Parameters
        ----------
        node_id   : topology node identifier
        node_type : e.g. "file_server", "web_server" — used to select templates
        base_url  : public base URL of the ShadowMesh backend,
                    e.g. "http://172.20.0.1:8000"
        count     : number of documents to generate (default 2)

        Returns
        -------
        List[DocMeta] — one entry per generated document
        """
        from backend.deception.canary import canary_manager

        # Select templates appropriate for this node type
        candidates = [t for t in _TEMPLATES if node_type in t["node_types"]]
        if not candidates:
            candidates = _DEFAULT_TEMPLATES

        # Pick `count` distinct templates (wrap if fewer candidates than count)
        selected: List[Dict] = []
        for i in range(count):
            selected.append(candidates[i % len(candidates)])

        results: List[DocMeta] = []
        for tmpl in selected:
            token_id = uuid.uuid4().hex
            canary_rel  = f"/api/canary/{token_id}"
            doc_rel     = f"/api/docs/{token_id}/{tmpl['filename']}"
            full_canary = f"{base_url.rstrip('/')}{canary_rel}"

            # Render document bytes
            try:
                data = self._render(tmpl, canary_rel, full_canary)
            except Exception as exc:
                log.error("[docgen] Render failed for %s (%s): %s",
                          tmpl["filename"], node_id, exc)
                continue

            fmt  = tmpl["fmt"]
            mime = _MIME[fmt]

            # Register canary token (reuses existing canary_manager)
            from backend.models import CanaryToken
            token = CanaryToken(
                token_id   = token_id,
                node_id    = node_id,
                token_url  = canary_rel,
                token_type = "document",
                label      = tmpl["filename"],
            )
            canary_manager._tokens[token_id] = token

            # Store document bytes
            self._store.put(_DocEntry(
                token_id   = token_id,
                node_id    = node_id,
                filename   = tmpl["filename"],
                mime       = mime,
                data       = data,
                created_at = time.time(),
            ))

            log.info("[docgen] Generated %s (%s, %d bytes) for node %s token %s",
                     tmpl["filename"], fmt, len(data), node_id, token_id)

            results.append(DocMeta(
                token_id   = token_id,
                node_id    = node_id,
                filename   = tmpl["filename"],
                fmt        = fmt,
                url        = doc_rel,
                canary_url = canary_rel,
            ))

        return results

    def _render(self, tmpl: Dict, canary_url: str, full_canary_url: str) -> bytes:
        fmt = tmpl["fmt"]
        if fmt == "txt":
            return _render_txt(tmpl["title"], canary_url, full_canary_url)
        if fmt == "docx":
            return _render_docx(tmpl["title"], tmpl["subject"], canary_url, full_canary_url)
        if fmt == "pdf":
            return _render_pdf(tmpl["title"], tmpl["subject"], canary_url, full_canary_url)
        raise ValueError(f"Unknown format: {fmt!r}")

    def get_document(self, token_id: str) -> Optional[_DocEntry]:
        """Retrieve a stored document by token_id. Returns None if missing or expired."""
        return self._store.get(token_id)

    def clear_for_node(self, node_id: str) -> None:
        """Remove all documents for a node (called on container teardown)."""
        removed = self._store.remove_for_node(node_id)
        if removed:
            log.debug("[docgen] Cleared %d document(s) for node %s", removed, node_id)

    def stats(self) -> Dict[str, int]:
        return {"stored_documents": self._store.count()}


# Module-level singleton
doc_generator = DocGenerator()
